from telethon import TelegramClient
from docx import Document
import docx
from docx.shared import Inches
import os
import asyncio
from datetime import datetime, timezone
import re

api_id = 25730961
api_hash = 'bc5faeb33eb16909b64fdc3bb397e3aa'

client = TelegramClient('test_tg', api_id, api_hash)
output_path = os.path.join(os.getcwd(), "Telegram_posts.docx")
download_path = 'downloads'
os.makedirs(download_path, exist_ok=True)
doc = Document()

channel_username = input("Введите название канала или ссылку: ")
user_date = input("Введите начальную дату для парсинга в формате ГГГГ-ММ-ДД: ")
try:
    start_date = datetime.strptime(
        user_date, "%Y-%m-%d").replace(tzinfo=timezone.utc)
except ValueError:
    print("Неверный формат даты. Используйте формат ГГГГ-ММ-ДД.")
    exit()


def add_hyperlink(paragraph, url, text, color, underline):
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    new_run = docx.oxml.shared.OxmlElement('w:r')

    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


async def main():
    await client.start()

    try:
        channel = await client.get_entity(channel_username)
    except Exception as e:
        print(f"Не удалось найти канал {channel_username}. Ошибка: {e}")
        await client.disconnect()
        exit()

    async for message in client.iter_messages(channel):
        print(message.date)

        if message.date < start_date:
            break

        doc.add_paragraph(f"Дата: {message.date.strftime(
            '%Y-%m-%d %H:%M:%S')}", 'Heading 3')

        if message.text:
            paragraph = doc.add_paragraph("Текст: ")
            text_parts = re.split(r'(https?://\S+)', message.text)

            for part in text_parts:
                if part.startswith("http"):
                    add_hyperlink(paragraph, part, part, 'FF8822', False)
                else:
                    paragraph.add_run(part)

        if message.photo:

            file_path = await message.download_media(file=download_path)
            if file_path and file_path.lower().endswith(('.png', '.jpg', '.jpeg')):

                try:
                    doc.add_picture(file_path, width=Inches(5))
                except Exception:
                    print("Ошибка при добавлении фото")
                os.remove(file_path)
            else:
                print(f"Пропущен неподдерживаемый файл: {file_path}")

        doc.add_paragraph("=" * 40)

    doc.save(output_path)
    print("Документ сохранен как Telegram_messages.docx")
    print(f"Документ сохранен по пути: {output_path}")

    await client.disconnect()

asyncio.run(main())
